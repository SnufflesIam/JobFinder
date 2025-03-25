import sqlite3, re, os, requests, webbrowser, time
from datetime import date
from bs4 import BeautifulSoup
try:
    from docx import Document
except ModuleNotFoundError as e:
    import pip
    print(f"{e}. Importing docx.")
    pip.main(['install', "docx"])
    from docx import Document
try:
    from selenium import webdriver
except ModuleNotFoundError as e:
    import pip
    print(f"{e}. Importing selenium.")
    pip.main(['install', "selenium"])
    from selenium import webdriver
from selenium.webdriver.common.by import By



#Saved all data into the database
#Takes 2 values, jobType must be either an 'A' for 'Apply' or 'I' for ignore.
#jobObject which contains all the details for each job listings
def SaveToDBS(job_Type, job_Object):
    jobType = job_Type
    jobObject = job_Object
    conn = sqlite3.connect('autoApply.db')
    today = date.today().strftime("%d %B %Y")
    try:
        if jobType == 'A':
            listing = jobObject['jobListing']
            owner = jobObject['listingOwner']
            role = jobObject['jobRole']
            conn.execute(f'''INSERT INTO ApplyFor(date, JobType, listingOwner, jobRole, URL)
    VALUES ("{today}","{jobType}","{owner}","{role}","{listing}")''')
        elif jobType == 'I':
            listing = jobObject['jobListing']
            ignore = jobObject['ignoreReason']
            conn.execute(f''' INSERT INTO IgnoreList (date, JobType, URL, ignoreReason)
    VALUES ("{today}","{jobType}","{listing}","{ignore}")
    ''')
        else:
            print("Invalid job type.")
        conn.commit()
        conn.close()
        
    except Exception as e:
        print(f"Line 52: There was an error inserting data into database. Error {e}")
        print(f"**jobType: {jobType}\njobObject {jobObject}")
        conn.commit()
        conn.close()

    
SEEKoptions = {
    "advertiser": '//*[@data-automation="advertiser-name"]',
    "job": '//h1[@data-automation="job-detail-title"]',
    "body": '//*[@data-automation="jobAdDetails"]',
    "URL": '//*[@data-automation="jobTitle"]'
}

GradConnectionOptions = {
    "advertiser": ['h1','employers-panel-title'],
    "job": ['h1','employers-profile-h1'],
    "body": ['div','campaign-content-container'] 
    }

IndeedOptions = {
    "advertiser": '//div[@data-testid="inlineHeader-companyName"]',
    "job": '//h1[@data-testid="jobsearch-JobInfoHeader-title"]',
    "body": '//div[@id="jobDescriptionText"]',
    "URL": '//h2/a[@class="jcs-JobTitle css-1baag51 eu4oa1w0"]'
    }

#This function filters out jobs that contain phrases or words that would normally cause me to ignore the job.
#The filtered jobs will still be saved, but they want be automatically applied for
#This function takes a list of URL's to be scrapes
#Returns 2 lists, a list of jobs that I should apply for, and a list of jobs that I should avoid
def FilterJobListings(jobList):
    ignorePhrases = [
        '[\d+]{1,3} years.{0,20}experience',
        '(final|last|second|penultimate).{0,20}year',
        'completed.{0,30}degree',
        '[Ss]enior .{0,20} [Dd]eveloper',
        '[Ss]enior.{0,20}[Ee]gineer',
        '[Ff]ulltime',
        '[Tt]echnical [Ll]ead',
        'CFO'
        ]
    ignoredJobs = []
    applyJobs = []
    jobListings = jobList
    SearchOptions = ""

    #Search each job listing page to see if it contains an phrase that would make it an undesireable job e.g. requires years of
    #experience in the role. Add the joh listing to the appropriate list.
    for listing in jobListings:
        match listing['website']:
            case "Seek":
                SearchOptions = SEEKoptions 
            case "GradConnection":
                SearchOptions = GradConnectionOptions
            case "Indeed":
                SearchOptions = IndeedOptions
        if listing['dynamic'] == False:
            PageToScrape = requests.get(listing['link'])
            soup = BeautifulSoup(PageToScrape.text, "html.parser")
            for phrase in ignorePhrases:
                elementType = SearchOptions['body'][1]
                if listing['website'] == 'Seek':
                    text = soup.find(f"{SearchOptions['body'][0]}", attrs={f"{SearchOptions['body'][1]}": f"{SearchOptions['body'][2]}"}).get_text()
                elif listing['website'] == 'GradConnection':
                    text = soup.find(f"{SearchOptions['body'][0]}", class_=f"{SearchOptions['body'][1]}").get_text()
                match = re.search(phrase, text)
                #I'm doing a quick check to break the loop if a match is found
                #Note I check this twice because i don't want to skip to the 'else' after every phrase in ignoredPhrases
                if match is not None:
                    break
            if match is not None:
                ignoredJobs.append({
                    "website": f"{listing['website']}",
                    "jobListing": f"{listing['link']}",
                    "ignoreReason": f"Found the phrase '{match.group()}'."
                    })
            else:
                if listing['website'] == 'Seek':
                    advertiserName = soup.find(f"{SearchOptions['advertiser'][0]}", attrs={f"{SearchOptions['advertiser'][1]}": f"{SearchOptions['advertiser'][2]}"}).get_text()
                    jobRole = soup.find(f"{SearchOptions['job'][0]}", attrs={f"{SearchOptions['job'][1]}": f"{SearchOptions['job'][2]}"}).get_text()
                elif listing['website'] == 'GradConnection':
                    advertiserName = soup.find(f"{SearchOptions['advertiser'][0]}", class_=f"{SearchOptions['advertiser'][1]}").get_text()
                    jobRole = soup.find(f"{SearchOptions['job'][0]}", class_=f"{SearchOptions['job'][1]}").get_text()
                applyJobs.append({
                    "website": f"{listing['website']}",
                    "jobListing": f"{listing['link']}",
                    "listingOwner": f"{advertiserName}",
                    "jobRole": f"{jobRole}"
                    })
        else:
            try:
                # set up Safari WebDriver using SafariDriverManager 
                driver = webdriver.Chrome()

                # open the specified URL in the browser 
                driver.get(listing['link'])
                
                #Get the job listings present on the web page
                body = driver.find_element(By.XPATH, SearchOptions['body']).text
                for phrase in ignorePhrases:
                    match = re.search(phrase, body)
                    if match is not None:
                        break
                if match is not None:
                    ignoredJobs.append({
                        "website": f"{listing['website']}",
                        "jobListing": f"{listing['link']}",
                        "ignoreReason": f"Found the phrase '{match.group()}'."
                        })
                else:
                    advertiserName = driver.find_element(By.XPATH, SearchOptions['advertiser']).text
                    jobRole = driver.find_element(By.XPATH, SearchOptions['job']).text
                    applyJobs.append({
                            "website": f"{listing['website']}",
                            "jobListing": f"{listing['link']}",
                            "listingOwner": f"{advertiserName}",
                            "jobRole": f"{jobRole}"
                            })
                # close the browser
                driver.quit()
            except Exception as e:
                print(f"Error: {e}\n**AN ERROR OCCURRED ACCESSING {listing['link']}\n***Try close all Chrome windows and try again.***\n**You may need to authenticate, they detected I'm a robot***\n")
                driver.quit()

    return (applyJobs, ignoredJobs)

#This function will take a generic coverletter, and insert listing-specific information into it
#It takes expects to recieve the following:
##document_name: the name of the generic cover letter
##save_path: the file path to save the edited cover letters
##listing: the object containing the listing, it's expected to contain the  website, recruiter and job role
#'recuiter', 'job-name' and the website the listing is sourced from.
def EditCoverletter(document_name, listing):
    #The word document is arranged using tables, so the text is accessed by getting the text from the cell at position 0,2
    document = Document(f"{document_name}")
    savePath = f"{os.getcwd()}/CoverLetters/"
    recruiter = listing['listingOwner']
    jobRole = listing['jobRole']
    website = listing['website']
    table = document.tables
    para_row = 0
    para_column = 2
    paragraphs = len(table[0].cell(para_row,para_column).paragraphs)
    today = date.today().strftime("%d %B %Y")
    
    #Insert the listing-specific information into the coverletter
    for para in range(paragraphs):
        newPara = table[0].cell(para_row,para_column).paragraphs[para].text
        newPara = newPara.replace("<RECRUITER>", f"{recruiter}")
        newPara = newPara.replace("<JOB ROLE>", f"{jobRole}")
        newPara = newPara.replace("<DATE>", f"{today}")
        newPara = newPara.replace("<WEBSITE>", f"{website}")
        #Omit details of my job for GradConnection jobs
        if website == "GradConnection":
            OGText = "I am an intelligence analyst for the ADF with a current TSPV,"
            newText = "I am an analyst for the ADF,"
            newPara = newPara.replace(OGText, newText)
        table[0].cell(para_row,para_column).paragraphs[para].text = newPara
            

    #Save the coverletters into a folder entitled today's date
    today = date.today().strftime("%Y%m%d")
    newDir = f"{today}/"
    path = os.path.join(savePath, newDir)
    if not os.path.exists(path):
        os.mkdir(path)
    try:
        document.save(f"{path}{recruiter}_coverletter.docx")
    except Exception as e:
        print(f"There was an error saving coverletter for the following listing:\n{listing}\nError:{e}")

#This function grabs all job listings available on a listing page
#It takes one value, the URL taken from the listing. It's expected the the URL is copied after the user has
#entered all their search criteria onto the jobsite (such as location, job titles, work hours etc.)
def GrabJobListings(website):
    
    #If the website's dynamic content loads without being interacted with, use BeautifulSoup
    if website['dynamic'] == False:
        #Set up the variables for the URL to scrape
        try:
            PageToScrape = requests.get(f"{website['link']}")
        except Exception as e:
            print(f"An error has occurred trying to connect to the website. Check the URL and ensure you are connected to the internet\n {e}")
            return
        soup = BeautifulSoup(PageToScrape.text, "html.parser")
        print(f"line 228: {soup}")


        #Find all the links in the URL
        string = ""
        REGEX = website['jobREGEX']
        for link in soup.find_all('a'):
            link = link.get('href')
            if link is not None:
                string += f"~{link}"
        #Filter all links to only the job IDs, remove doubles and add to a jobs list
        REGEX = website['jobREGEX']
        jobIDs = re.findall(REGEX, string)
        jobIDs = list(dict.fromkeys(jobIDs))
        jobLinks = []
        for item in jobIDs:
            jobLinks.append({
                "website": f"{website['website']}",
                "dynamic": False,
                "broadSearch": False,
                "link" : f"{website['baseURL']}{item}"
                })
    #If the website's dynamic content doesn't load until it is interacted with, used Selenium
    else:
        #Establish the relative path for the website to locate joblistings
        relativePath = ""
        match website['website']:
            case "Indeed":
                relativePath = IndeedOptions['URL']
                # set up Safari WebDriver using SafariDriverManager
                #options = get_default_chrome_options()
                driver = webdriver.Chrome()
                
                # open the specified URL in the browser 
                driver.get(website['link'])

                #Get the job listings present on the web page
                jobIDs = driver.find_elements(By.XPATH, relativePath)

                #Add all the found listings to the jobLinks list
                jobLinks = []
                for item in jobIDs:
                    jobLinks.append({
                        "website": f"{website['website']}",
                        "dynamic": True,
                        "broadSearch": False,
                        "link": f"{item.get_attribute('href')}"
                        })

                # close the browser
                driver.quit()
            case "Seek":
                relativePath = SEEKoptions['URL']
                # set up Safari WebDriver using SafariDriverManager
                #options = get_default_chrome_options()
                driver = webdriver.Chrome()
                
                # open the specified URL in the browser 
                driver.get(website['link'])


                #Get the job listings present on the web page
                jobIDs = driver.find_elements(By.XPATH, relativePath)

                #Add all the found listings to the jobLinks list
                jobLinks = []
                for item in jobIDs:
                    jobLinks.append({
                        "website": f"{website['website']}",
                        "dynamic": True,
                        "broadSearch": False,
                        "link": f"{item.get_attribute('href')}"
                        })
                # close the browser
                driver.quit()
            
        
    return jobLinks

#Function opens the link
#Expects to recieve a string with the letters denoting which websites to open and an object with the listings
def OpenJob(site_options, listings):
    siteOptions = site_options
    listings = listings
    for listing in listings:
        if 'a' in siteOptions and len(siteOptions) == 1:
            webbrowser.open(listing['jobListing'])
        elif 's' in siteOptions and listing['website'] == "Seek":
            webbrowser.open(listing['jobListing'])
        elif 'g' in siteOptions and listing['website'] == "GradConnection":
            webbrowser.open(listing['jobListing'])
        elif 'i' in siteOptions and listing['website'] == "Indeed":
            webbrowser.open(listing['jobListing'])
        elif 'o' in siteOptions and listing['broadSearch'] == True:
            webbrowser.open(listing['jobListing'])



#findJobs uses search phrases and posts them to apiJobs (https://app.apijobs.dev)
#searchTerms - string of terms to search for, e,g 'software egineer'
def searchByTerms(searchTerms, country):
    #Establish headers to connect to apiJobs
    headers = {
        'apikey': 'bebce871b2d5057209648407ba4163babe35396b6b7a68790e3a4da0f7a87d29',
        'Content-Type': 'application/json',
    }

    #Post each phrase to apiJob
    searchTerms = re.findall('[\w]+', searchTerms)
    Results = []
    for item in searchTerms:
        data = '{"sort_by":"created_at","q":'+f'"{item}"'+',"country":'+f'"{country}"'+'}'
        response = requests.post('https://api.apijobs.dev/v1/job/search', headers=headers, data=data)
        Results.append(response.json())
        
    #Check if each listing a phrase to ignore in the description
    #If it doesn't contain an 'ignore phrase', open the website
    searchResults = []
    ignoreList = [
        "[\d +]+ years"
        ]
    #Iterate each item in the results returned by the api
    #Save the title, description, url and if it's ignored, save the reason
    try:
        for item in Results[0]['hits']:
            for phrase in ignoreList:
                ig = re.findall(phrase, item['description'])
                if len(ig) < 1:
                    searchResults.append({
                        "website": item['website'],
                        "jobListing": item['url'],
                        "listingOwner":item['hiring_organization_name'],
                        "jobRole": item['title'],
                        "description": item['description'],
                        "broadSearch": True,
                        "ignored": False,
                        "ignoreReason": ""
                        })
                else:
                    searchResults.append({
                        "website": item['website'],
                        "jobListing": item['url'],
                        "listingOwner":item['hiring_organization_name'],
                        "jobRole": item['title'],
                        "description": item['description'],
                        "broadSearch": True,
                        "ignored": True,
                        "ignoreReason": f"Found the phrase: {ig}"
                        })
    except Exception as e:
        print(f"Error iterating through broad search results: {e}\nResults[1]['hits'] contained: {Results[1]['hits']}")

    return(searchResults)












    
